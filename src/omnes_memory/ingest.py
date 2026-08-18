"""Инжест документов: txt/md/pdf/docx с мягкими зависимостями (фаза 4.1)."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

log = logging.getLogger("omnes.ingest")

SUPPORTED = {".txt", ".md", ".pdf", ".docx"}


def read_text_file(path: Path) -> str:
    """UTF-8 с фолбэком на cp1251 (типично для старых русских документов)."""
    raw = path.read_bytes()
    for encoding in ("utf-8", "cp1251"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader  # noqa: PLC0415 — опциональная зависимость
    except ImportError as e:
        raise RuntimeError("pypdf не установлен: pip install '.[docs]'") from e
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def read_docx(path: Path) -> str:
    try:
        import docx  # noqa: PLC0415 — опциональная зависимость
    except ImportError as e:
        raise RuntimeError("python-docx не установлен: pip install '.[docs]'") from e
    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def read_document(path: str | Path) -> tuple[str, dict[str, Any]]:
    """Читает файл по расширению; возвращает (текст, метаданные)."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"файл не найден: {p}")
    suffix = p.suffix.lower()
    if suffix not in SUPPORTED:
        raise ValueError(f"не поддерживается {suffix!r} (поддерживаются: {sorted(SUPPORTED)})")
    if suffix == ".pdf":
        text = read_pdf(p)
    elif suffix == ".docx":
        text = read_docx(p)
    else:
        text = read_text_file(p)
    meta = {"file_name": p.name, "size_bytes": p.stat().st_size, "format": suffix}
    return text, meta
